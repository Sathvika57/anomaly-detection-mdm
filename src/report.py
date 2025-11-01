# src/report.py
"""
Summarize anomalies, create filtered alert CSV, optionally send email, and generate simple PDF/DOCX report.
"""
import os
import pandas as pd
from datetime import datetime
from src.policy_filters import postfilter_alerts
from src.alerts import send_email_if_needed

ANOM_DIR = "data/data_anomalies"
REPORT_DIR = "data/reports"
os.makedirs(REPORT_DIR, exist_ok=True)

def latest_anomaly_file():
    files = sorted([f for f in os.listdir(ANOM_DIR) if f.endswith(".csv")])
    if not files:
        return None
    return os.path.join(ANOM_DIR, files[-1])

def generate_pdf_summary(summary_text: str, out_path: str):
    try:
        from fpdf import FPDF
    except Exception:
        print("FPDF not installed, skipping PDF creation.")
        return
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in summary_text.split("\n"):
        pdf.multi_cell(0, 8, line)
    pdf.output(out_path)
    print("Saved PDF report:", out_path)

def generate_docx_summary(summary_text: str, out_path: str):
    try:
        from docx import Document
    except Exception:
        print("python-docx not installed, skipping DOCX creation.")
        return
    doc = Document()
    doc.add_heading("MDM Anomaly Detection Report", 0)
    for para in summary_text.split("\n"):
        doc.add_paragraph(para)
    doc.save(out_path)
    print("Saved DOCX report:", out_path)

def main():
    anom = latest_anomaly_file()
    if not anom:
        print("No anomaly file found. Run detect_anomaly.py first.")
        return
    df = pd.read_csv(anom, parse_dates=["window_start"])
    total = len(df)
    anomalies = int(df["is_anomaly"].sum())
    unique_devices = int(df["device_id"].nunique()) if "device_id" in df.columns else 0

    # filtered alerts
    alerts = postfilter_alerts(df[df["is_anomaly"]==True], top_k=200)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    alerts_file = os.path.join(REPORT_DIR, f"alerts_{ts}.csv")
    alerts.to_csv(alerts_file, index=False)
    print("Saved alerts CSV:", alerts_file)

    # summary text
    summary_text = f"MDM Anomaly Detection Report\nRun time (UTC): {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}\n"
    summary_text += f"Input anomaly file: {os.path.basename(anom)}\nTotal records processed: {total}\nTotal anomalies flagged: {anomalies}\nUnique devices: {unique_devices}\nAlerts saved to: {alerts_file}\n"
    summary_text += "\nTop alerts preview:\n"
    if not alerts.empty:
        summary_text += alerts.head(10).to_string(index=False)
    else:
        summary_text += "No actionable alerts after filtering."

    # save summary CSV
    summary_csv = os.path.join(REPORT_DIR, f"report_summary_{ts}.csv")
    df.describe(include="all").to_csv(summary_csv)
    print("Saved descriptive summary:", summary_csv)

    # PDF and DOCX
    pdf_out = os.path.join(REPORT_DIR, f"anomaly_report_{ts}.pdf")
    docx_out = os.path.join(REPORT_DIR, f"anomaly_report_{ts}.docx")
    generate_pdf_summary(summary_text, pdf_out)
    generate_docx_summary(summary_text, docx_out)

    # send email if alerts exist (environment-controlled)
    if not alerts.empty:
        send_email_if_needed(subject=f"MDM Alerts: {len(alerts)} detected", body=summary_text, attachment_path=alerts_file)

if __name__ == "__main__":
    main()
